from datetime import date
from typing import List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from your_utils import (
    _compute_column_widths,
    format_date,
    add_header,
    add_footer,
    insert_table_of_contents,
    build_executive_summary
)
from table_utils import style_table, add_mini_table_to_cell
from html_parser import clean_html_and_extract_tables
from http_client import ProjectMeta, Issue
from html2docx import html2docx


class ReportBuilder:
    def __init__(self, meta: ProjectMeta, issues: List[Issue], report_date: date):
        self.meta = meta
        self.issues = issues
        self.report_date = report_date
        self.doc = Document()

    def init_styles(self):
        sect = self.doc.sections[0]
        sect.page_width, sect.page_height = Inches(11), Inches(8.5)
        for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(sect, m, Inches(0.5))

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        for lvl, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
            h = styles[lvl]
            h.font.name = "Calibri"
            h.font.size = Pt(size)

    def build_cover_page(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Regional Issues Report\n")
        run.font.size = Pt(28)

        for txt in (
            "Mini Group / Eleven Degrees Consulting",
            f"Project: {self.meta.name}",
            f"Region: {self.meta.region}",
            f"Date: {format_date(self.report_date)}",
        ):
            self.doc.add_paragraph(txt).alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

    def build_toc_and_summary(self):
        insert_table_of_contents(self.doc, levels=2)
        self.doc.add_page_break()
        build_executive_summary(self.doc, self.issues)

    def _region_key(self, raw_region: Any) -> str:
        if isinstance(raw_region, (list, tuple)):
            for v in raw_region:
                if isinstance(v, str) and v.strip():
                    return v.strip()
            return "UNKNOWN"
        if isinstance(raw_region, str) and raw_region.strip():
            return raw_region.strip()
        return "UNKNOWN"

    def build_sections(self):
        groups: dict[str, List[Issue]] = {}

        for issue in self.issues:
            key = self._region_key(issue.region)
            groups.setdefault(key, []).append(issue)

        for idx, (branch, branch_issues) in enumerate(groups.items()):
            if idx > 0:
                self.doc.add_section(WD_SECTION.NEW_PAGE)

            add_header(
                self.doc,
                left_text="minigroup_logo.png",
                center_text=f"{self.meta.name} – {branch}",
                right_text="minigroup.png"
            )

            self.doc.add_heading(f"Branch: {branch}", level=1)

            for label, val in [
                ("Start Date", self.meta.start_date),
                ("Status", self.meta.status),
                ("Branch Manager", self.meta.bm),
                ("Operations Manager", self.meta.om),
                ("Supervisor", self.meta.sup),
            ]:
                if val:
                    p = self.doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(str(val))
            self.doc.add_paragraph()

            for issue in branch_issues:
                self._render_issue(issue)

    def _render_issue(self, issue: Issue):
        self.doc.add_heading(f"Issue: {issue.title}", level=2)

        # 🔍 Log raw HTML presence of <table>
        print(f"\n[DEBUG] Issue '{issue.title}': Raw HTML contains table?")
        for label, html in [
            ("Description", issue.description_html),
            ("Recommendation", issue.recommendation or ""),
            ("Mgmt Comment 1", issue.mgmt_comment1 or ""),
            ("Mgmt Comment 2", issue.mgmt_comment2 or ""),
            ("Additional Table", issue.table_html or "")
        ]:
            contains_table = "<table" in html.lower()
            print(f"  - {label}: {'✅ Table found' if contains_table else '❌ No table'}")

        # ✅ Use raw HTML tables only for Description
        desc_text, desc_tables = clean_html_and_extract_tables(issue.description_html, preserve_tables_as_html=True)
        recomm_text, recomm_tables = clean_html_and_extract_tables(issue.recommendation or "")
        mgmt1_text, mgmt1_tables = clean_html_and_extract_tables(issue.mgmt_comment1 or "")
        mgmt2_text, mgmt2_tables = clean_html_and_extract_tables(issue.mgmt_comment2 or "")
        table_text, table_tables = clean_html_and_extract_tables(issue.table_html or "")
        implication = issue.implication or ""
        cost = f"${issue.cost_impact:,.2f}"

        print(f"[DEBUG] Issue '{issue.title}':")
        print(f"  - Description tables: {len(desc_tables)}")
        print(f"  - Recommendation tables: {len(recomm_tables)}")
        print(f"  - Mgmt Comment 1 tables: {len(mgmt1_tables)}")
        print(f"  - Mgmt Comment 2 tables: {len(mgmt2_tables)}")
        print(f"  - Additional tables: {len(table_tables)}")

        fields = [
            ("Issue Title", issue.title, []),
            ("Severity", issue.severity.capitalize(), []),
            ("Description", desc_text, desc_tables),  # raw HTML tables
            ("Implication", implication, []),
            ("Cost Impact", cost, []),
            ("Mgmt Comment 1", mgmt1_text, mgmt1_tables),
            ("Mgmt Comment 2", mgmt2_text, mgmt2_tables),
            ("Recommendation", recomm_text, recomm_tables),
            ("Additional Table", table_text, table_tables),
        ]

        # Build table in Word
        tbl = self.doc.add_table(rows=len(fields), cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = False

        widths = _compute_column_widths(
            [[label, text] for label, text, _ in fields],
            max_total_width_inches=Inches(10).inches
        )
        tbl.columns[0].width = widths[0]
        tbl.columns[1].width = widths[1]

        for i, (label, text, tables) in enumerate(fields):
            left, right = tbl.rows[i].cells

            # Label cell
            left.text = label
            if left.paragraphs and left.paragraphs[0].runs:
                left.paragraphs[0].runs[0].bold = True

            # Clear right cell
            for p in list(right.paragraphs):
                p._element.getparent().remove(p._element)

            # Add main text
            right.add_paragraph(text or "")

            # ⬇️ Handle tables for Description field differently
            if label == "Description":
                for raw_html in tables:
                    temp_doc = Document()
                    html2docx(raw_html, temp_doc)
                    for para in temp_doc.paragraphs:
                        right._element.append(para._element)
                    for tbl in temp_doc.tables:
                        right._element.append(tbl._element)
            else:
                # Normal parsed tables
                for headers, rows in tables:
                    print(f"    ➤ Rendering table in '{label}' with headers: {headers} and {len(rows)} rows")
                    add_mini_table_to_cell(right, headers, rows)

        self.doc.add_paragraph()  # spacing after issue



    def build(self, output_path=None):
        self.init_styles()
        self.build_cover_page()
        self.build_toc_and_summary()
        self.build_sections()
        #self.build_footer()
        if output_path:
            self.doc.save(output_path)
            print(f"✅ Report saved to {output_path}")
        else:
            self.save()

    def build_footer(self):
        # Add footer with region and page number
        self.doc.sections[0].footer.is_linked_to_previous = False
        self.doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add left and right footer text
        self.doc.sections[0].footer.paragraphs[0].add_run(
            f"Regional Issues Report – {self.meta.region} | "
        ).bold = True
        self.doc.sections[0].footer.paragraphs[0].add_run("Page ")
        self.doc.sections[0].footer.paragraphs[0].add_run().add_field("PAGE")
 