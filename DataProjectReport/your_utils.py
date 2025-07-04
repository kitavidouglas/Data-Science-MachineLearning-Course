from typing import Any, List
from datetime import date
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import uuid
import os

# ─── 1. Safe Conversion ───────────────────────────────
def ensure_str(val: Any, max_len: int = None) -> str:
    if val is None:
        return ''
    if isinstance(val, (bytes, bytearray)):
        try:
            s = val.decode('utf-8')
        except UnicodeDecodeError:
            s = val.decode('latin-1', errors='ignore')
    elif isinstance(val, str):
        s = val.strip()
    elif isinstance(val, (int, float)):
        s = str(val)
    elif isinstance(val, list):
        s = ', '.join(map(str, val))
    elif isinstance(val, dict):
        s = ', '.join(f"{k}={v}" for k, v in val.items())
    else:
        try:
            s = str(val).strip()
        except Exception:
            s = ''
    if max_len and len(s) > max_len:
        return s[:max_len - 3] + '...'
    return s

# ─── 2. Date Formatter ───────────────────────────────
def format_date(d: date, fmt: str = "%B %d, %Y") -> str:
    return d.strftime(fmt)

# ─── 3. Add Header with Logos + Metadata ─────────────
def add_header(doc: Document, left_text: str, center_text: str, right_text: str):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        table = header.add_table(rows=1, cols=3, width=Inches(11))
        table.autofit = False
        widths = [Inches(3), Inches(5), Inches(3)]

        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width

        left, center, right = table.rows[0].cells

        try:
            left.paragraphs[0].add_run().add_picture(left_text, width=Inches(1.2))
        except Exception:
            left.paragraphs[0].add_run(left_text).bold = True

        center_p = center.paragraphs[0]
        center_p.add_run(center_text).bold = True
        center_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        try:
            right.paragraphs[0].add_run().add_picture(right_text, width=Inches(1.2))
        except Exception:
            right.paragraphs[0].add_run(right_text).bold = True
        right.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ─── 4. Add Footer with Page Number ───────────────────
def add_footer(doc: Document, left: str, right: str, page_number: bool = True):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        table = footer.add_table(rows=1, cols=3, width=Inches(11))
        table.autofit = False
        widths = [Inches(3), Inches(5), Inches(3)]

        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width

        left_cell, center_cell, right_cell = table.rows[0].cells
        left_cell.paragraphs[0].add_run(left).italic = True
        center_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if page_number:
            p = right_cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = p.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
            run.font.size = Pt(9)

# ─── 5. Table of Contents Field ───────────────────────
def insert_table_of_contents(doc: Document, levels: int = 2):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# ─── 6. Insert Chart from Matplotlib ──────────────────
def insert_chart(fig, doc: Document, caption: str = ""):
    chart_path = f"/tmp/chart_{uuid.uuid4().hex}.png"
    fig.savefig(chart_path, bbox_inches="tight", dpi=300)
    doc.add_picture(chart_path, width=Inches(6.5))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].italic = True
    os.remove(chart_path)

# ─── 7. Executive Summary Generator ───────────────────
def build_executive_summary(doc: Document, issues):
    doc.add_heading("Executive Summary", level=1)

    severity_counts = Counter(issue.severity for issue in issues if issue.severity)
    total_cost = sum(issue.cost_impact for issue in issues if issue.cost_impact)

    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = para.add_run(f"This report covers {len(issues)} issue(s) reported in the selected region.\n")
    run.font.size = Pt(11)

    for sev in ["high", "medium", "low"]:
        count = severity_counts.get(sev, 0)
        para.add_run(f" - {sev.capitalize()} severity: {count} issue(s)\n")

    para.add_run(f"\nEstimated total cost impact: ${total_cost:,.2f}\n")
    doc.add_paragraph()  # extra spacing

# ─── 8. Column Width Computation for Report Tables ────
def _compute_column_widths(
    text_matrix: List[List[Any]],
    max_total_width_inches: float = 10.0,
    min_width_inches: float = 0.5
) -> List[Any]:
    if not text_matrix:
        return []

    cols = list(zip(*text_matrix))
    scores = []
    total_score = 0.0

    for col in cols:
        lengths = [len(str(val)) for val in col]
        max_len = max(lengths)
        is_sentence = any(len(str(val).split()) > 5 for val in col)
        is_numeric = all(str(val).replace('.', '', 1).isdigit() for val in col if val)
        weight = 1.5 if is_sentence else 0.5 if is_numeric else 1.0
        score = max_len * weight
        scores.append(score)
        total_score += score

    total_score = total_score or 1.0
    raw_widths = [(s / total_score) * max_total_width_inches for s in scores]
    capped = [max(min_width_inches, w) for w in raw_widths]

    if sum(capped) > max_total_width_inches:
        factor = max_total_width_inches / sum(capped)
        capped = [w * factor for w in capped]

    return [Inches(w) for w in capped]

# Expose as public API
compute_column_widths = _compute_column_widths
