from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.table import _Cell
from typing import Tuple, List
import re

from typing import List, TypedDict, Union, Optional, Tuple

# --- Typed definitions for structured blocks ---
class TextRun(TypedDict):
    text: str
    bold: bool
    italic: bool
    href: Optional[str]

class ParagraphBlock(TypedDict):
    type: str  # "paragraph"
    runs: List[TextRun]

class TableBlock(TypedDict):
    type: str  # "table"
    headers: List[str]
    rows: List[List[str]]

Block = Union[ParagraphBlock, TableBlock]

# ✅ MAIN: parse HTML → structured blocks
def parse_html_content(html: str) -> List[Block]:
    from bs4 import BeautifulSoup, NavigableString, Tag
    soup = BeautifulSoup(html or "", "html.parser")

    # Remove scripts/styles and normalize line breaks
    for tag in soup(["script", "style"]): tag.decompose()
    for br in soup.find_all("br"): br.replace_with("\n")

    blocks: List[Block] = []

    # 🔧 Use soup.contents instead of soup.body to ensure full parsing
    elements = soup.contents  # <-- This is the fix

    for elem in elements:
        if isinstance(elem, Tag) and elem.name == "table":
            # Handle table
            headers = [th.get_text(strip=True) for th in elem.find_all("th")]
            rows = []
            for tr in elem.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            blocks.append({"type": "table", "headers": headers, "rows": rows})
        else:
            # Handle paragraphs with inline formatting
            text = elem.get_text(separator="\n", strip=True)
            if not text:
                continue
            for para in text.split("\n\n"):
                runs: List[TextRun] = []

                def recurse(node, bold=False, italic=False, href=None):
                    if isinstance(node, NavigableString):
                        txt = node.strip()
                        if txt:
                            runs.append({"text": txt, "bold": bold, "italic": italic, "href": href})
                    elif isinstance(node, Tag):
                        b = bold or node.name in ("strong", "b")
                        i = italic or node.name in ("em", "i")
                        h = href or (node.get("href") if node.name == "a" else None)
                        for c in node.contents:
                            recurse(c, b, i, h)

                recurse(BeautifulSoup(para, "html.parser"))
                if runs:
                    blocks.append({"type": "paragraph", "runs": runs})

    return blocks

# ✅ DOCX render: use container for inline cells!
def render_html_to_docx(doc: Document, html_str: str, container: Optional[_Cell] = None):
    blocks = parse_html_content(html_str)

    for block in blocks:
        parent = container if container else doc

        if block["type"] == "paragraph":
            p = parent.add_paragraph()
            for run in block["runs"]:
                r = p.add_run(run["text"])
                if run["bold"]:
                    r.bold = True
                if run["italic"]:
                    r.italic = True
                if run.get("href"):
                    r.font.underline = True

        elif block["type"] == "table":
            headers = block["headers"]
            rows = block["rows"]
            cols = len(headers) or (len(rows[0]) if rows else 0)

            table = parent.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            table.autofit = True

            if headers:
                hdr_cells = table.rows[0].cells
                for idx, text in enumerate(headers):
                    hdr_cells[idx].text = text

            for i, row_data in enumerate(rows, start=1):
                cells = table.rows[i].cells
                for j, cell_text in enumerate(row_data):
                    cells[j].text = cell_text

    if not container:
        doc.add_paragraph()  # Spacer for top-level

# ✅ CLEAN: return plain text & tables for safe CSV or nested insertion


def clean_html_and_extract_tables(html: str, preserve_tables_as_html=False):
    """
    Extracts clean text and tables from an HTML string.

    If preserve_tables_as_html is True:
        - Instead of parsing tables into headers/rows,
        - The tables are extracted as raw HTML blocks for later rendering.

    Returns:
        (text, tables)
        text: main body text as plain text
        tables: list of either (headers, rows) or raw HTML strings
    """
    if not html:
        return "", []

    soup = BeautifulSoup(html, "html.parser")

    tables = []
    for table in soup.find_all("table"):
        if preserve_tables_as_html:
            tables.append(str(table))  # Keep raw HTML table
        else:
            # ✅ Convert table to (headers, rows)
            headers = []
            rows = []

            first_row = table.find("tr")
            if first_row:
                headers = [th.get_text(strip=True) for th in first_row.find_all(["th", "td"])]

            for tr in table.find_all("tr")[1:]:
                row = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                rows.append(row)

            tables.append((headers, rows))

        table.decompose()  # remove the table so it doesn’t show in text

    # Extract remaining text without tables
    text = soup.get_text(separator=" ", strip=True)
    return text, tables

def extract_faux_table(text: str):
    lines = re.split(r"[•·•\n]", text)
    rows = []
    for line in lines:
        parts = re.split(r" {2,}|\t", line.strip())  # multiple spaces or tab
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 2:
            rows.append(parts)
    return rows