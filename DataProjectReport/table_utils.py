from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from typing import List

# ─── Style a Table with Caption + Equal Column Widths ─────────────
def style_table(table, caption: str = "", total_width: float = 10.0):
    """
    Apply 'Table Grid' style, insert optional caption above,
    and distribute `total_width` inches equally across all table columns.   
    """
    # 1) Optional caption above the table
    if caption:
        tbl_elm = table._element
        p_elm   = OxmlElement("w:p")
        r_elm   = OxmlElement("w:r")
        t_elm   = OxmlElement("w:t")
        t_elm.text = caption
        r_elm.append(t_elm)
        p_elm.append(r_elm)
        tbl_elm.addprevious(p_elm)

    # 2) Apply style and disable default autofit
    table.style   = "Table Grid"
    table.autofit = False

    # 3) Distribute total width evenly
    col_count = len(table.columns)
    if col_count > 0:
        width_per_col = Inches(total_width / col_count)
        for col in table.columns:
            for cell in col.cells:
                cell.width = width_per_col

    # 4) Center-align all cell text by default
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# ─── Add a Nested Mini Table Inside a Cell ────────────────────────
def add_mini_table_to_cell(cell, headers, rows):
    """
    Adds a well-styled mini-table to a Word cell.
    Handles missing headers/rows safely and formats properly.
    """
    if not headers and not rows:
        return

    # Add a spacer paragraph before the table
    cell.add_paragraph()

    # Determine column count
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    table = cell.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False

    # Equal width assignment (optional tuning)
    col_width = Inches(5.0 / col_count) if col_count else Inches(5.0)
    for col in table.columns:
        for c in col.cells:
            c.width = col_width

    # ✅ Add header row (bold and centered)
    if headers:
        hdr_cells = table.add_row().cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ✅ Add data rows (left-aligned)
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            p = row_cells[i].paragraphs[0]
            p.add_run(str(val))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT